"""OpenAI-compatible API layer: streaming, thinking mapping, image encoding."""
import base64
import io
import json
import os
from dataclasses import dataclass
from typing import AsyncIterator

import httpx
from PIL import Image
from pathlib import Path

MAX_IMAGE_EDGE = 1568  # OpenAI 视觉模型图片最大边

# 思考档位 -> 请求参数（provider 配置 thinking_extra 可按模型覆盖）
DEFAULT_THINKING = {
    "off": {},
    "low": {"reasoning_effort": "low"},
    "medium": {"reasoning_effort": "medium"},
    "high": {"reasoning_effort": "high"},
}


@dataclass
class Msg:
    role: str
    parts: list  # [{"type":"text","text":...}, {"type":"image_url",...}]
    id: int | None = None
    hidden: bool = False
    summary: bool = False


@dataclass
class StreamResult:
    text: str = ""
    reasoning: str = ""
    usage: dict | None = None
    error: str | None = None


# 附件文本提取支持的类型（图片走 encode_image，不在此列）
FILE_EXTS = (".txt", ".md", ".markdown", ".log", ".csv", ".json", ".py", ".js", ".ts",
             ".html", ".htm", ".css", ".xml", ".yaml", ".yml", ".toml", ".ini", ".cfg",
             ".sh", ".bat", ".cmd", ".ps1", ".tex", ".rst", ".pdf", ".docx", ".xlsx")
MAX_FILE_CHARS = 100_000  # 单个附件提取后截断长度


def extract_text(path: str) -> str:
    """把附件内容提取为纯文本。缺依赖时抛 RuntimeError，由调用方提示安装。"""
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise RuntimeError("读取 PDF 需要 pypdf：pip install pypdf")
        reader = PdfReader(str(p))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if ext == ".docx":
        try:
            import docx
        except ImportError:
            raise RuntimeError("读取 docx 需要 python-docx：pip install python-docx")
        d = docx.Document(str(p))
        lines = [para.text for para in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                lines.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(lines)
    if ext == ".xlsx":
        try:
            import openpyxl
        except ImportError:
            raise RuntimeError("读取 xlsx 需要 openpyxl：pip install openpyxl")
        wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
        out = []
        try:
            for ws in wb.worksheets:
                out.append(f"# Sheet: {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if v is None else str(v) for v in row]
                    if any(cells):
                        out.append(" | ".join(cells))
        finally:
            wb.close()
        return "\n".join(out)
    # 其余类型按纯文本读取，自动尝试常见编码
    for enc in ("utf-8", "gb18030", "latin-1"):
        try:
            return p.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return p.read_text(encoding="utf-8", errors="replace")


def encode_image(path: str, max_edge: int = MAX_IMAGE_EDGE) -> str:
    """压缩图片为 JPEG data URL。"""
    with Image.open(path) as im:
        im = im.convert("RGB")
        im.thumbnail((max_edge, max_edge))
        buf = io.BytesIO()
        im.save(buf, format="JPEG", quality=85)
    return "data:image/jpeg;base64," + base64.b64encode(buf.getvalue()).decode()


def parts_from_inputs(text: str, image_paths: list[str]) -> list:
    parts = [{"type": "text", "text": text}] if text else []
    for p in image_paths:
        parts.append({"type": "image_url", "image_url": {"url": encode_image(p)}})
    return parts or [{"type": "text", "text": ""}]


def to_api_messages(msgs: list[Msg], system_prompt: str) -> list[dict]:
    out = []
    if system_prompt:
        out.append({"role": "system", "content": system_prompt})
    for m in msgs:
        out.append({"role": m.role, "content": m.parts})
    return out


def thinking_params(cfg: dict, model: str, level: str) -> dict:
    params = dict(DEFAULT_THINKING.get(level, {}))
    extra = cfg.get("thinking_extra", {})
    if extra and model in extra:
        params.update(extra[model])
    return params


def api_key(cfg: dict) -> str:
    if cfg.get("api_key"):  # 配置文件内明文 key（/setting 写入）
        return cfg["api_key"]
    env = cfg.get("api_key_env", "")
    return os.environ.get(env, "") if env else ""


async def stream_chat(cfg: dict, model: str, messages: list[dict], thinking: str = "off", max_tokens: int | None = None) -> AsyncIterator[StreamResult]:
    """流式对话，逐块产出累积结果，最后一块带 usage/error。"""
    base = cfg["base_url"].rstrip("/")
    body = {
        "model": model,
        "messages": messages,
        "stream": True,
        "stream_options": {"include_usage": True},
    }
    body.update(thinking_params(cfg, model, thinking))
    if max_tokens:
        body["max_tokens"] = max_tokens
    headers = {}
    key = api_key(cfg)
    if key:
        headers["Authorization"] = f"Bearer {key}"
    acc = StreamResult()
    try:
        async with httpx.AsyncClient(timeout=httpx.Timeout(120.0, connect=20.0)) as client:
            async with client.stream("POST", f"{base}/chat/completions", json=body, headers=headers) as resp:
                if resp.status_code != 200:
                    acc.error = f"HTTP {resp.status_code}: {(await resp.aread()).decode(errors='replace')[:500]}"
                    yield acc
                    return
                async for line in resp.aiter_lines():
                    if not line.startswith("data:"):
                        continue
                    data = line[5:].strip()
                    if data == "[DONE]":
                        continue
                    try:
                        chunk = json.loads(data)
                    except json.JSONDecodeError:
                        continue
                    if chunk.get("error"):
                        acc.error = f"API error: {chunk['error']}"
                        yield acc
                        return
                    choices = chunk.get("choices") or []
                    delta = choices[0].get("delta", {}) if choices else {}
                    acc.text += delta.get("content") or ""
                    acc.reasoning += delta.get("reasoning_content") or ""
                    if chunk.get("usage"):
                        acc.usage = chunk["usage"]
                    yield acc
    except httpx.HTTPError as e:
        acc.error = f"网络错误 ({type(e).__name__}): {e or '连接中断或超时'}"
        yield acc
    except Exception as e:
        acc.error = f"请求异常 ({type(e).__name__}): {e}"
        yield acc
